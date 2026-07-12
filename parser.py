import fitz  # PyMuPDF,读 PDF 的库

# 支持的代码文件扩展名(代码本质是文本,直接读;分块时按函数/逻辑块切)
CODE_EXTS = (".c", ".h", ".cpp", ".hpp", ".cc", ".s", ".S", ".py",
            ".js", ".ts", ".java", ".rs", ".go", ".sh", ".lua")


def parse_pdf(file_path: str, mineru="auto", token: str = None) -> str:
    """解析 PDF。
    mineru=True:强制 MinerU 在线 API;False:强制 PyMuPDF;
    'auto'(默认):先 PyMuPDF(快),若文字极少(扫描件)自动转 MinerU OCR。
    MinerU 失败自动降级 PyMuPDF,不中断。"""
    if mineru is True:
        try:
            return parse_pdf_mineru(file_path, token=token)
        except Exception as e:
            print(f"[!] MinerU 解析失败({e}),降级 PyMuPDF")
    doc = fitz.open(file_path)
    pages = [p.get_text() for p in doc]
    doc.close()
    text = "\n".join(pages)
    # auto:平均每页文字太少 → 判定扫描件 → 自动转 MinerU OCR
    if mineru == "auto" and _looks_like_scanned(pages):
        print(f"[i] 检测到扫描件(每页文字极少),自动转 MinerU OCR...")
        try:
            return parse_pdf_mineru(file_path, token=token)
        except Exception as e:
            print(f"[!] MinerU 失败({e}),保留 PyMuPDF 结果")
    return text


def _looks_like_scanned(pages: list, threshold: int = 50) -> bool:
    """判断是否扫描件/图片型 PDF:平均每页非空字符 < threshold(默认 50)。
    文字型 PDF 每页通常几百字符,扫描件 PyMuPDF 提不出文字(接近 0)。"""
    if not pages:
        return False
    avg = sum(len(p.strip()) for p in pages) / len(pages)
    return avg < threshold


# ============ MinerU 在线 API(可选:扫描件 / 复杂版式 / 表格)============
def parse_pdf_mineru(file_path: str, token: str = None, language: str = None, timeout: int = 300):
    """用 MinerU 在线 API 解析 PDF。
    有 token(MINERU_TOKEN)→ 精准 API(200MB/200页,zip 含 markdown+json);
    无 token → Agent 轻量 API(10MB/20页,免登录,仅 markdown)。
    失败抛异常,由 parse_pdf 捕获后降级 PyMuPDF。"""
    import config
    if language is None: language = config.getd("mineru.language")
    file_name = file_path.replace("\\", "/").split("/")[-1]
    if token:
        return _mineru_precise(file_path, file_name, token, language, timeout)
    return _mineru_lite(file_path, file_name, language, timeout)


def _mineru_lite(file_path, file_name, language, timeout):
    """Agent 轻量 API(免 token):申请上传 → PUT → 轮询 → 取 markdown"""
    import requests, time
    r = requests.post("https://mineru.net/api/v1/agent/parse/file",
                      json={"file_name": file_name, "language": language, "is_ocr": True}, timeout=30)
    if r.status_code != 200 or r.json().get("code") != 0:
        raise RuntimeError(f"申请上传失败:{r.text[:120]}")
    d = r.json()["data"]
    task_id, file_url = d["task_id"], d["file_url"]
    with open(file_path, "rb") as f:
        up = requests.put(file_url, data=f, timeout=120)
    if up.status_code not in (200, 201):
        raise RuntimeError(f"上传失败 HTTP {up.status_code}")
    start = time.time()
    while time.time() - start < timeout:
        rr = requests.get(f"https://mineru.net/api/v1/agent/parse/{task_id}", timeout=30).json()
        state = rr["data"]["state"]
        if state == "done":
            return requests.get(rr["data"]["markdown_url"], timeout=30).text
        if state == "failed":
            raise RuntimeError(rr["data"].get("err_msg", "未知错误"))
        time.sleep(3)
    raise RuntimeError(f"轮询超时({timeout}s)")


def _mineru_precise(file_path, file_name, token, language, timeout):
    """精准 API(需 token):file-urls/batch 申请上传 → PUT → 轮询 → 解 zip 取 full.md"""
    import requests, time, zipfile, io
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    r = requests.post("https://mineru.net/api/v4/file-urls/batch",
                      headers=headers,
                      json={"files": [{"name": file_name, "is_ocr": True}],
                            "model_version": "vlm", "language": language}, timeout=30)
    if r.status_code != 200 or r.json().get("code") != 0:
        raise RuntimeError(f"申请上传失败:{r.text[:120]}")
    d = r.json()["data"]
    batch_id, urls = d["batch_id"], d["file_urls"]
    with open(file_path, "rb") as f:
        requests.put(urls[0], data=f, timeout=120)
    start = time.time()
    while time.time() - start < timeout:
        rr = requests.get(f"https://mineru.net/api/v4/extract-results/batch/{batch_id}",
                          headers=headers, timeout=30).json()
        item = rr["data"]["extract_result"][0]
        state = item["state"]
        if state == "done":
            zip_bytes = requests.get(item["full_zip_url"], timeout=120).content
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
                for name in z.namelist():
                    if name.endswith("full.md"):
                        return z.read(name).decode("utf-8")
            raise RuntimeError("结果 zip 里没找到 full.md")
        if state == "failed":
            raise RuntimeError(item.get("err_msg", "未知错误"))
        time.sleep(5)
    raise RuntimeError(f"轮询超时({timeout}s)")


def parse_text(file_path: str) -> str:
    """读取纯文本文件(.txt / .md),直接返回内容"""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def parse_docx(file_path: str) -> str:
    """读取 Word 文档(.docx),按"段落 + 表格"提取文字。
    注意:python-docx 只读新版 .docx,不读老的二进制 .doc。"""
    from docx import Document
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_epub(file_path: str) -> str:
    """读取 EPUB 电子书(.epub)。EPUB 本质是"一堆网页打包",
    每个章节是一个 html,用 BeautifulSoup 把正文提出来。"""
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    book = epub.read_epub(file_path)
    parts = []
    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        text = soup.get_text(separator="\n")
        clean = "\n".join(line for line in text.splitlines() if line.strip())
        if clean:
            parts.append(clean)
    return "\n\n".join(parts)


def parse_html(file_path: str) -> str:
    """读取网页(.html / .htm),先删掉脚本/样式/导航这些干扰,再提取正文。"""
    from bs4 import BeautifulSoup
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        html = f.read()
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    return "\n".join(line for line in text.splitlines() if line.strip())


def parse(file_path: str, mineru="auto", token: str = None) -> str:
    """根据文件后缀名,自动决定用哪种方式解析(总调度)。
    mineru:True=强制 MinerU;False=强制 PyMuPDF;'auto'(默认)=PDF 自动检测扫描件。"""
    lower = file_path.lower()
    if lower.endswith(".pdf"):
        return parse_pdf(file_path, mineru=mineru, token=token)
    elif lower.endswith((".txt", ".md")):
        return parse_text(file_path)
    elif lower.endswith(".docx"):
        return parse_docx(file_path)
    elif lower.endswith(".epub"):
        return parse_epub(file_path)
    elif lower.endswith((".html", ".htm")):
        return parse_html(file_path)
    elif lower.endswith(CODE_EXTS):       # 代码文件(.c/.py/.s...)本质是文本
        return parse_text(file_path)
    else:
        raise ValueError(f"暂不支持这种格式: {file_path}")


def is_code_file(file_path: str) -> bool:
    """判断是否代码文件(供 chunker 按函数/逻辑块切分用)"""
    return file_path.lower().endswith(CODE_EXTS)
